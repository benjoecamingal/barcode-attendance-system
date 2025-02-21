from flask import Flask, request, jsonify, session, send_file
from flask_cors import CORS
import pandas as pd
import os
import random
import string
from db_config import get_db_connection  # Import MySQL database connection
from barcode_generator import generate_barcode
from docx import Document
from docx.shared import Pt
from barcode import Code128
from barcode.writer import ImageWriter
import time
import io
from datetime import datetime, timedelta
import openpyxl
from openpyxl.styles import NamedStyle
from openpyxl.utils import get_column_letter

from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document

from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
import os

import tempfile

# Flask App Configuration
app = Flask(__name__)
app.secret_key = "your_secret_key"  # Required for session management
app.config['SESSION_COOKIE_NAME'] = 'my_session'  # Ensures session persistence
CORS(app, supports_credentials=True)  # Allow session cookies in CORS requests

UPLOAD_FOLDER = tempfile.gettempdir()  # Temporary directory for uploads
PROCESSED_FOLDER = tempfile.gettempdir()  # Temporary directory for processing


# Login Route
@app.route("/login", methods=["POST"])
def login():
    """Handles user login"""
    data = request.json
    username = data.get("username")
    password = data.get("password")

    db = get_db_connection()
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE username = %s AND password = %s", (username, password))
    user = cursor.fetchone()
    db.close()

    if user:
        session["user"] = username  # Store session
        session.permanent = True  # Ensure session persists
        print("Session after login:", session)  # Debugging
        return jsonify({"message": "Login successful"}), 200
    else:
        return jsonify({"error": "Invalid credentials"}), 401

def generate_barcode_image(barcode_number, output_path):
    try:
        # Remove .png extension if it exists
        if output_path.endswith('.png'):
            output_path = output_path[:-4]
            
        # Create barcode
        barcode = Code128(str(barcode_number), writer=ImageWriter())
        
        # Make sure directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Save barcode - python-barcode will add .png extension automatically
        saved_path = barcode.save(output_path)
        print(f"Successfully saved barcode to: {saved_path}")
        
        return True
    except Exception as e:
        print(f"Error generating barcode: {str(e)}")
        return False



def generate_word_document(dataframe, output_path, barcode_paths):
    doc = Document()

    section = doc.sections[0]
    section.page_height = Inches(11)  # Letter size height
    section.page_width = Inches(8.5)  # Letter size width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    student_count = 0

    for index, row in dataframe.iterrows():
        name = row.get("Name", "Unknown")
        position = row.get("Position", "Unknown")
        department = row.get("Department", "Unknown")
        barcode_number = row.get("Barcode", "")

        if not barcode_number:
            continue

        barcode_path = barcode_paths.get(barcode_number)

        if not barcode_path or not os.path.exists(barcode_path):
            continue

        # Add Name (Centered, Large, Bold)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
        run = paragraph.add_run(f"{name}\n")
        run.bold = True
        run.font.size = Pt(20)  # Set font size to 20

        # Add Position (Centered)
        p_pos = doc.add_paragraph(f"Position: {position}")
        p_pos.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add Department (Centered)
        p_dept = doc.add_paragraph(f"Department: {department}")
        p_dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add Barcode Image (Centered)
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run_img = p_img.add_run()
            run_img.add_picture(barcode_path, width=Inches(2.5))  # Adjust width to be readable
        except Exception as e:
            error_p = doc.add_paragraph("[ERROR: Could not add barcode image]")
            error_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add Separator Line (Centered)
        p_line = doc.add_paragraph("-" * 30)  # Create separator line
        p_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Ensure it's centered

        # Two students per page
        student_count += 1
        if student_count % 2 == 0:
            doc.add_page_break()

    try:
        doc.save(output_path)
    except Exception as e:
        print(f"Error saving document: {str(e)}")


def generate_unique_barcode(cursor):
    """Generate a unique barcode ensuring no duplicates in the database."""
    while True:
        barcode = generate_barcode()
        cursor.execute("SELECT COUNT(*) FROM students WHERE barcode = %s", (barcode,))
        if cursor.fetchone()[0] == 0:
            return barcode

# File Upload & Barcode Processing
@app.route("/upload", methods=["POST"])
def upload_file():
    temp_file = None
    output_docx = None
    barcode_paths = {}
    
    try:
        if "file" not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
            
        file = request.files["file"]
        if file.filename == "":
            return jsonify({"error": "No selected file"}), 400

        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
        file.save(temp_file.name)
        
        # Close the file handle explicitly
        temp_file.close()
        
        # Read Excel file
        df = pd.read_excel(temp_file.name, engine="openpyxl")
        
        # Basic data validation
        df.dropna(how="all", inplace=True)
        df.fillna("", inplace=True)
        
        if df.empty:
            raise ValueError("Uploaded file is empty")
            
        # Verify required columns
        required_columns = ["Name", "Batch", "Position", "Department", "School"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            raise ValueError(f"Missing required columns: {', '.join(missing_columns)}")

        # Connect to database
        db = get_db_connection()
        cursor = db.cursor()
        
        try:
            # Generate barcodes for each row if needed
            if "Barcode" not in df.columns:
                df["Barcode"] = df.apply(lambda _: generate_unique_barcode(cursor), axis=1)

            # Store barcode images
            for _, row in df.iterrows():
                barcode_number = row["Barcode"]
                base_path = os.path.join(tempfile.gettempdir(), f"barcode_{barcode_number}")
                barcode_path = f"{base_path}.png"
                
                if generate_barcode_image(barcode_number, base_path):
                    barcode_paths[barcode_number] = barcode_path

            # Insert data into database
            for _, row in df.iterrows():
                student_dict = row.to_dict()
                columns = ", ".join(student_dict.keys())
                placeholders = ", ".join(["%s"] * len(student_dict))
                values = tuple(student_dict.values())
                query = f"INSERT INTO students ({columns}) VALUES ({placeholders})"
                cursor.execute(query, values)
            
            db.commit()

            # Generate Word document
            output_docx = os.path.join(tempfile.gettempdir(), "student_barcodes.docx")
            generate_word_document(df, output_docx, barcode_paths)

            # Read the file into memory before sending
            with open(output_docx, 'rb') as f:
                file_data = io.BytesIO(f.read())

            # Clean up files before sending response
            cleanup_files(temp_file.name, output_docx, barcode_paths)

            # Send file from memory
            file_data.seek(0)
            return send_file(
                file_data,
                as_attachment=True,
                download_name="student_barcodes.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as db_error:
            db.rollback()
            raise db_error

        finally:
            cursor.close()
            db.close()

    except Exception as e:
        # Clean up files in case of error
        cleanup_files(
            temp_file.name if temp_file else None,
            output_docx if output_docx else None,
            barcode_paths
        )
        return jsonify({"error": str(e)}), 500


def cleanup_files(temp_file_path, output_docx_path, barcode_paths):
    """Helper function to clean up temporary files"""
    # Clean up temporary Excel file
    if temp_file_path and os.path.exists(temp_file_path):
        try:
            os.unlink(temp_file_path)
        except (PermissionError, OSError):
            pass  # Ignore errors if file is still in use
            
    # Clean up generated Word document  
    if output_docx_path and os.path.exists(output_docx_path):
        try:
            os.unlink(output_docx_path)
        except (PermissionError, OSError):
            pass
            
    # Clean up barcode images
    for path in barcode_paths.values():
        if os.path.exists(path):
            try:
                os.unlink(path)
            except (PermissionError, OSError):
                pass

    

@app.route("/add_student", methods=["POST"])
def add_student():
    """Handles adding an individual student to the database and generates Word document"""
    temp_file = None
    output_docx = None
    barcode_paths = {}
    
    try:
        if "user" not in session:
            return jsonify({"error": "Unauthorized"}), 403

        if not request.is_json:
            return jsonify({"error": "Invalid request format. JSON required"}), 400

        data = request.get_json()
        required_fields = ["name", "batch", "position", "department", "school"]
        missing_fields = [field for field in required_fields if not data.get(field)]

        if missing_fields:
            return jsonify({"error": f"Missing required fields: {', '.join(missing_fields)}"}), 400

        # Create a single-row DataFrame from the submitted data
        df = pd.DataFrame([{
            'Name': data["name"].strip(),
            'Batch': data["batch"].strip(),
            'Position': data["position"].strip(),
            'Department': data["department"].strip(),
            'School': data["school"].strip()
        }])

        # Basic data validation
        if df.empty:
            raise ValueError("No valid data provided")

        # Connect to database
        db = get_db_connection()
        cursor = db.cursor()
        
        try:
            # Generate unique barcode
            barcode = generate_unique_barcode(cursor)
            df["Barcode"] = barcode

            # Generate and store barcode image
            base_path = os.path.join(tempfile.gettempdir(), f"barcode_{barcode}")
            barcode_path = f"{base_path}.png"
            
            if generate_barcode_image(barcode, base_path):
                barcode_paths[barcode] = barcode_path
            else:
                raise Exception("Failed to generate barcode image")

            # Insert data into database
            student_dict = df.iloc[0].to_dict()
            columns = ", ".join(student_dict.keys())
            placeholders = ", ".join(["%s"] * len(student_dict))
            values = tuple(student_dict.values())
            query = f"INSERT INTO students ({columns}) VALUES ({placeholders})"
            cursor.execute(query, values)
            
            db.commit()

            # Generate Word document
            output_docx = os.path.join(tempfile.gettempdir(), f"student_barcode_{barcode}.docx")
            generate_word_document(df, output_docx, barcode_paths)

            # Read the file into memory before sending
            with open(output_docx, 'rb') as f:
                file_data = io.BytesIO(f.read())

            # Clean up files before sending response
            cleanup_files(None, output_docx, barcode_paths)

            # Send file from memory
            file_data.seek(0)
            return send_file(
                file_data,
                as_attachment=True,
                download_name=f"student_barcode_{data['name']}.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as db_error:
            db.rollback()
            raise db_error

        finally:
            cursor.close()
            db.close()

    except Exception as e:
        # Clean up files in case of error
        cleanup_files(
            temp_file.name if temp_file else None,
            output_docx if output_docx else None,
            barcode_paths
        )
        return jsonify({"error": str(e)}), 500

# Logout Route
@app.route("/logout", methods=["POST"])
def logout():
    """Handles user logout"""
    session.pop("user", None)  # Remove user from session
    return jsonify({"message": "Logged out successfully"}), 200


@app.route('/scan', methods=['POST'])
def process_scan():
    """Handles barcode scanning requests from PyQt GUI"""
    data = request.get_json()
    barcode = data.get("barcode")

    if not barcode:
        return jsonify({"success": False, "message": "No barcode received"}), 400

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Check if student exists
        cursor.execute("SELECT id, name, department FROM students WHERE barcode = %s", (barcode,))
        student = cursor.fetchone()

        if not student:
            cursor.close()
            conn.close()
            return jsonify({"success": False, "message": "Student not found"}), 404

        student_id, student_name, department = student

        # Clear any unread results before executing the next query
        cursor.fetchall()

        # Check if the student already has a time_in for today
        cursor.execute(
            "SELECT id, time_in, time_out FROM attendance WHERE student_id = %s AND date = CURDATE()",
            (student_id,)
        )
        record = cursor.fetchone()

        if record:
            attendance_id, time_in, time_out = record

            # Clear any unread results before the next query
            cursor.fetchall()

            if time_in and not time_out:
                # Second scan → Time Out
                cursor.execute(
                    "UPDATE attendance SET time_out = CURRENT_TIME WHERE id = %s",
                    (attendance_id,)
                )
                conn.commit()
                cursor.close()
                conn.close()
                return jsonify({
                    "success": True,
                    "message": f"✅ Time Out recorded for {student_name}",
                    "name": student_name,
                    "department": department,
                    "time": time_out,
                    "status": "Time Out",
                    "date": "Today"
                })
            else:
                cursor.close()
                conn.close()
                return jsonify({
                    "success": False,
                    "message": "❌ Already Timed Out for Today"
                })
        else:
            # First scan → Time In
            cursor.execute(
                "INSERT INTO attendance (student_id, date, time_in) VALUES (%s, CURDATE(), CURRENT_TIME)",
                (student_id,)
            )
            conn.commit()
            cursor.close()
            conn.close()
            return jsonify({
                "success": True,
                "message": f"✅ Time In recorded for {student_name}",
                "name": student_name,
                "department": department,
                "time": "Now",
                "status": "Time In",
                "date": "Today"
            })

    except Exception as e:
        cursor.close()
        conn.close()
        return jsonify({"success": False, "message": str(e)}), 500
    

@app.route('/attendance', methods=['GET'])
def get_attendance():
    """Fetch attendance records with optional filters"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        # Get filters from request
        batch = request.args.get('batch')
        position = request.args.get('position')
        department = request.args.get('department')
        date = request.args.get('date')  # Format: YYYY-MM-DD

        # Start SQL Query
        query = """
            SELECT s.name, s.batch, s.position, s.department, 
                   a.date, a.time_in, a.time_out
            FROM attendance a
            JOIN students s ON a.student_id = s.id
            WHERE 1=1
        """
        filters = []

        # Add filters dynamically
        if batch:
            query += " AND s.batch = %s"
            filters.append(batch)
        if position:
            query += " AND s.position = %s"
            filters.append(position)
        if department:
            query += " AND s.department = %s"
            filters.append(department)
        if date:
            query += " AND a.date = %s"
            filters.append(date)

        query += " ORDER BY a.date DESC"

        # Execute the query with filters
        cursor.execute(query, tuple(filters))
        records = cursor.fetchall()

        # Convert time fields to string
        for record in records:
            record["time_in"] = str(record["time_in"]) if record["time_in"] else "N/A"
            record["time_out"] = str(record["time_out"]) if record["time_out"] else "N/A"

        return jsonify(records)

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        cursor.close()
        conn.close()


@app.route('/filters', methods=['GET'])
def get_filters():
    """Fetch distinct values for Batch, Position, and Department from database"""
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)

        cursor.execute("SELECT DISTINCT batch FROM students WHERE batch IS NOT NULL")
        batches = [row["batch"] for row in cursor.fetchall()]

        cursor.execute("SELECT DISTINCT position FROM students WHERE position IS NOT NULL")
        positions = [row["position"] for row in cursor.fetchall()]

        cursor.execute("SELECT DISTINCT department FROM students WHERE department IS NOT NULL")
        departments = [row["department"] for row in cursor.fetchall()]

        return jsonify({
            "batches": batches,
            "positions": positions,
            "departments": departments
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        cursor.close()
        db.close()


def format_excel_time(worksheet):
    """Apply time formatting to Time In and Time Out columns"""
    # Create time style
    time_style = NamedStyle(name='time_style')
    time_style.number_format = 'HH:MM:SS'
    
    # Find Time In and Time Out columns
    headers = [cell.value for cell in worksheet[1]]
    time_in_col = headers.index('Time In') + 1 if 'Time In' in headers else None
    time_out_col = headers.index('Time Out') + 1 if 'Time Out' in headers else None
    
    # Apply formatting
    if time_in_col:
        for cell in worksheet[get_column_letter(time_in_col)][1:]:
            cell.style = time_style
    if time_out_col:
        for cell in worksheet[get_column_letter(time_out_col)][1:]:
            cell.style = time_style

@app.route('/attendance/download', methods=['GET'])
def download_attendance():
    """Generate and send attendance data as Excel file"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        # Get filters from request
        batch = request.args.get('batch')
        position = request.args.get('position')
        department = request.args.get('department')
        date = request.args.get('date')

        # Base query
        query = """
            SELECT 
                s.name as 'Name',
                s.batch as 'Batch',
                s.position as 'Position',
                s.department as 'Department',
                a.date as 'Date',
                a.time_in as 'Time In',
                a.time_out as 'Time Out'
            FROM attendance a
            JOIN students s ON a.student_id = s.id
            WHERE 1=1
        """
        filters = []

        # Add filters
        if batch:
            query += " AND s.batch = %s"
            filters.append(batch)
        if position:
            query += " AND s.position = %s"
            filters.append(position)
        if department:
            query += " AND s.department = %s"
            filters.append(department)
        if date:
            query += " AND a.date = %s"
            filters.append(date)

        query += " ORDER BY a.date DESC, s.name ASC"

        # Execute query
        cursor.execute(query, tuple(filters))
        records = cursor.fetchall()
        print("Fetched Records:", records)  # Debugging

        # ✅ Convert timedelta (Time In/Out) to HH:MM:SS string
        for record in records:
            if isinstance(record['Time In'], timedelta):  # ✅ Use timedelta directly
                record['Time In'] = str(record['Time In'])
            if isinstance(record['Time Out'], timedelta):  # ✅ Use timedelta directly
                record['Time Out'] = str(record['Time Out'])

        # Convert to DataFrame
        df = pd.DataFrame(records)

        # Create Excel file in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Attendance', index=False)

            # Auto-adjust column widths
            worksheet = writer.sheets['Attendance']
            for idx, col in enumerate(df.columns):
                max_length = max(
                    df[col].astype(str).apply(len).max(),
                    len(col)
                ) + 2
                worksheet.column_dimensions[chr(65 + idx)].width = max_length

        output.seek(0)

        # Generate filename with current date
        current_date = datetime.now().strftime('%Y%m%d')
        filename = f"attendance_report_{current_date}.xlsx"

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        print("Error in download_attendance():", str(e))  # Debugging
        return jsonify({"error": str(e)}), 500

    finally:
        cursor.close()
        conn.close()




# Run Flask App
if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5000, debug=True)
